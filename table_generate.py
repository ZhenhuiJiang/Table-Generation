"""
randomly generate tables and bodies to detect table

table start with @ , end with *
body start with %, end with $
large head start with #, end with $
small head start with %, end with $
"""

table_start='@'
table_end='*'
body_start='%'
body_end='$'
large_head_start='#'
large_head_end='$'
small_head_start='%'
small_head_end='$'


from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Length
from docx.shared import RGBColor
from docx.shared import Inches
from random import Random
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from random import Random
from faker import Faker
from docx.enum.table import WD_ALIGN_VERTICAL

fake = Faker()

random = Random()


def random_str(randomlength=8):
    str = ''
    chars = 'AaBbCcDdEeFfGgHhIiJjKkLlMmNnOoPpQqRrSsTtUuVvWwXxYyZz0123456789                 '
    length = len(chars) - 1
    for i in range(randomlength):
        str += chars[random.randint(0, length)]
    return str


def random_num(randomlength=8):
    str = ''
    chars = '0123456789'
    length = len(chars) - 1
    for i in range(randomlength):
        str += chars[random.randint(0, length)]
    return str


def set_A4_size(document):
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)


def body(row=3, left_indent=0, right_indent=0, line_spacing=18, space_before=0, space_after=0, qianzhui=True):
    no_char = row * 100
    #        words=fake.text()[:no_char]
    words = fake.text()[:no_char]
    if qianzhui:
        #        words=fake.text()[:no_char]
        words = body_start + words[1:-1] + body_end
    paragraph = document.add_paragraph(words)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(left_indent)
    paragraph_format.right_indent = Inches(right_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)


def body2(row=3, left_indent=0, right_indent=0, line_spacing=18, space_before=0, space_after=0, qianzhui=True):
    no_char = row * 100
    #        words=fake.text()[:no_char]
    words = random_str(no_char)
    if qianzhui:
        #        words=fake.text()[:no_char]
        words = body_start + words[1:-1] + body_end
    paragraph = document.add_paragraph(words)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(left_indent)
    paragraph_format.right_indent = Inches(right_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)


def head(size=20, bold=True, italic=False, max_num_str=10, min_num_str=1, xiao_qianzhui=True, da_qianzhui=False):
    no_char = random.randint(min_num_str, max_num_str)
    words = fake.text()[:no_char]
    if xiao_qianzhui:
        words = small_head_start + words[1:-1] + small_head_end
    if da_qianzhui:
        words = large_head_start + words[1:-1] + large_head_end
    run = document.add_paragraph()
    head = run.add_run(words)
    # 颜色
    head.font.color.rgb = RGBColor(0, 0, 0)
    font = head.font
    # 字体，大小
    font.name = 'Calibri'
    font.size = Pt(size)
    # 加粗
    font.bold = bold
    font.italic = italic


def table(rows=5, cols=3, right_align_start=1, min_num_str=5, max_num_str=13,
          min_word_str=5, max_word_str=18,
          first_rows_no_number=0, first_rows_no_desc=0, first_rows_no_desc_col_start=0,
          num_of_cols_start_with_num=1, qianzhui=True, height=0.6, midrow_no_number=0):
    if max_num_str > 13:
        return ("ERROR,max_num_str should be less than 13 ")
    if first_rows_no_number > rows:
        return ("ERROR:first_rows_no_number should be less than rows")
    if first_rows_no_desc > rows:
        return ("ERROR:first_rows_no_desc should be less than rows")
    if first_rows_no_desc_col_start > cols:
        return ("ERROR:first_rows_no_desc_col_start should be less than cols")
    # table = document.add_table(rows = 1,cols = 3, style='Table Grid')
    table = document.add_table(rows=rows, cols=cols, style=None)

    #    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True

    # 前几行没有数字只有description
    if first_rows_no_number > 0:
        for row in range(first_rows_no_number):
            # 随机生成最大n个字符
            hdr_cells = table.rows[row].cells
            # 为每一个单元格赋值
            # 注：值都要为字符串类型
            n = random.randint(min_word_str, max_word_str)
            words = fake.text()[:n]
            hdr_cells[0].paragraphs[0].add_run(words).bold = True
        for i in range(right_align_start, cols):
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    # 前几行没有description只有数字
    if first_rows_no_desc > 0:
        for row in range(first_rows_no_desc):
            # 随机生成最大n个字符
            hdr_cells = table.rows[row].cells
            for col in range(first_rows_no_desc_col_start, cols):
                n = random.randint(min_num_str, max_num_str)
                words = random_num(n)
                hdr_cells[col].text = words
            for i in range(right_align_start, cols):
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    whole_rows = max(first_rows_no_number, first_rows_no_desc)
    # 剩下的都有的
    for row in range(whole_rows, rows):
        # 随机生成最大n个字符
        hdr_cells = table.rows[row].cells
        # 为每一个单元格赋值
        # 注：值都要为字符串类型
        for col in range(num_of_cols_start_with_num):
            n = random.randint(min_word_str, max_word_str)
            words = fake.text()[:n]
            hdr_cells[col].text = words
        for col in range(num_of_cols_start_with_num, cols):
            n = random.randint(min_num_str, max_num_str)
            words = random_num(n)
            hdr_cells[col].text = words
        for i in range(right_align_start, cols):
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    if midrow_no_number > 0:
        hdr_cells = table.rows[midrow_no_number].cells
        n = random.randint(min_word_str, max_word_str)
        words = fake.text()[:n]
        for col in range(cols):
            hdr_cells[col].text = ' '
        hdr_cells[0].paragraphs[0].add_run(words)
    if qianzhui:
        table.cell(0, 0).text = table_start
        table.cell(rows - 1, cols - 1).text = table_end
        table.cell(rows - 1, cols - 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        table.cell(rows - 1, cols - 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT.EXACTLY
        row.height = Cm(height)


def table_first_two_merge(rows=5, cols=4, right_align_start=2, min_num_str=5, max_num_str=13,
                          min_word_str=5, max_word_str=18,
                          first_rows_no_number=0, first_rows_no_desc=0, first_rows_no_desc_col_start=0,
                          num_of_cols_start_with_num=2, qianzhui=True, height=0.6, midrow_no_number=0):
    if max_num_str > 13:
        return ("ERROR,max_num_str should be less than 13 ")
    if first_rows_no_number > rows:
        return ("ERROR:first_rows_no_number should be less than rows")
    if first_rows_no_desc > rows:
        return ("ERROR:first_rows_no_desc should be less than rows")
    if first_rows_no_desc_col_start > cols:
        return ("ERROR:first_rows_no_desc_col_start should be less than cols")
    # table = document.add_table(rows = 1,cols = 3, style='Table Grid')
    table = document.add_table(rows=rows, cols=cols, style=None)
    #    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True

    # 前几行没有数字只有description
    if first_rows_no_number > 0:
        for row in range(first_rows_no_number):
            # 随机生成最大n个字符
            hdr_cells = table.rows[row].cells
            # 为每一个单元格赋值
            # 注：值都要为字符串类型
            n = random.randint(min_word_str, max_word_str)
            words = fake.text()[:n]
            hdr_cells[0].merge(hdr_cells[1])
            hdr_cells[0].paragraphs[0].add_run(words).bold = True
        for i in range(right_align_start, cols):
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    # 前几行没有description只有数字
    if first_rows_no_desc > 0:
        for row in range(first_rows_no_desc):
            # 随机生成最大n个字符
            hdr_cells = table.rows[row].cells
            for col in range(first_rows_no_desc_col_start, cols):
                n = random.randint(min_num_str, max_num_str)
                words = random_num(n)
                hdr_cells[col].text = words
            for i in range(right_align_start, cols):
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    whole_rows = max(first_rows_no_number, first_rows_no_desc)
    # 剩下的都有的
    for row in range(whole_rows, rows):
        # 随机生成最大n个字符
        hdr_cells = table.rows[row].cells
        hdr_cells[0].merge(hdr_cells[1])
        # 为每一个单元格赋值
        # 注：值都要为字符串类型
        for col in range(num_of_cols_start_with_num):
            n = random.randint(min_word_str, max_word_str)
            words = fake.text()[:n]
            hdr_cells[col].text = words
        for col in range(num_of_cols_start_with_num, cols):
            n = random.randint(min_num_str, max_num_str)
            words = random_num(n)
            hdr_cells[col].text = words
        for i in range(right_align_start, cols):
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    if midrow_no_number > 0:
        hdr_cells = table.rows[midrow_no_number].cells
        n = random.randint(min_word_str, max_word_str)
        words = fake.text()[:n]
        for col in range(cols):
            hdr_cells[col].text = ' '
        hdr_cells[0].paragraphs[0].add_run(words)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT.EXACTLY
        row.height = Cm(height)
    if qianzhui:
        table.cell(0, 0).text = table_start
        table.cell(rows - 1, cols - 1).text = table_end
        table.cell(rows - 1, cols - 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        table.cell(rows - 1, cols - 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


document = Document()
func = (

    lambda: table(rows=random.randint(8, 10), cols=random.randrange(3, 5), qianzhui=True,
                  min_num_str=2, max_num_str=10, midrow_no_number=random.randrange(3, 7)),
    lambda: table(rows=random.randint(2, 6), cols=random.randrange(3, 5)
                  , qianzhui=True, first_rows_no_number=random.randrange(0, 3)),
    lambda: table(rows=random.randint(4, 8), cols=random.randrange(3, 5)
                  , first_rows_no_desc=random.randrange(0, 3)
                  , qianzhui=True, first_rows_no_desc_col_start=random.randint(0, 3)),
    lambda: body2(row=random.randint(3, 10), qianzhui=True),
    lambda: body2(row=random.randint(12, 30), qianzhui=True),
    lambda: body(row=random.randint(6, 16), qianzhui=True),
    lambda: body2(row=random.randint(6, 16), qianzhui=True),
    lambda: head(size=20, bold=True, italic=False, max_num_str=20, min_num_str=10, da_qianzhui=True,
                 xiao_qianzhui=False),
    lambda: head(size=12, bold=True, italic=False, max_num_str=30, min_num_str=20, da_qianzhui=False,
                 xiao_qianzhui=True),
    lambda: table_first_two_merge(rows=random.randrange(4, 10),
                                  cols=random.randrange(3, 5),
                                  first_rows_no_number=random.randrange(0, 3),
                                  first_rows_no_desc=random.randrange(0, 3),
                                  first_rows_no_desc_col_start=random.randrange(2, 3)
                                  , qianzhui=True, midrow_no_number=random.randrange(2, 4))
)

fake.seed(1)
random.seed(1)
for _ in range(3000):
    random.choice(func)()
document.save('qianzhui_a.docx')

print(1)

document = Document()
fake.seed(2)
random.seed(2)
for _ in range(3000):
    random.choice(func)()

document.save('qianzhui_b.docx')

print(1)

fake.seed(3)
random.seed(3)
for _ in range(1500):
    random.choice(func)()
document.save('qianzhui_c.docx')

print(1)

document = Document()
fake.seed(4)
random.seed(4)
for _ in range(3000):
    random.choice(func)()

document.save('qianzhui_d.docx')

print(1)

fake.seed(5)
random.seed(5)
for _ in range(1500):
    random.choice(func)()
document.save('qianzhui_e.docx')

print(1)

document = Document()
fake.seed(6)
random.seed(6)
for _ in range(3000):
    random.choice(func)()

document.save('qianzhui_f.docx')







document = Document()
func = (

    lambda: table(rows=random.randint(8, 10), cols=random.randrange(3, 5), qianzhui=False,
                  min_num_str=2, max_num_str=10, midrow_no_number=random.randrange(3, 7)),
    lambda: table(rows=random.randint(2, 6), cols=random.randrange(3, 5)
                  , qianzhui=False, first_rows_no_number=random.randrange(0, 3)),
    lambda: table(rows=random.randint(4, 8), cols=random.randrange(3, 5)
                  , first_rows_no_desc=random.randrange(0, 3)
                  , qianzhui=False, first_rows_no_desc_col_start=random.randint(0, 3)),
    lambda: body2(row=random.randint(3, 10), qianzhui=True),
    lambda: body2(row=random.randint(12, 30), qianzhui=True),
    lambda: body(row=random.randint(6, 16), qianzhui=True),
    lambda: body2(row=random.randint(6, 16), qianzhui=True),
    lambda: head(size=20, bold=True, italic=False, max_num_str=20, min_num_str=10, da_qianzhui=True,
                 xiao_qianzhui=False),
    lambda: head(size=12, bold=True, italic=False, max_num_str=30, min_num_str=20, da_qianzhui=False,
                 xiao_qianzhui=True),
    lambda: table_first_two_merge(rows=random.randrange(4, 10),
                                  cols=random.randrange(3, 5),
                                  first_rows_no_number=random.randrange(0, 3),
                                  first_rows_no_desc=random.randrange(0, 3),
                                  first_rows_no_desc_col_start=random.randrange(2, 3)
                                  , qianzhui=False, midrow_no_number=random.randrange(2, 4))
)

fake.seed(1)
random.seed(1)
for _ in range(3000):
    random.choice(func)()
document.save('no_qianzhui_a.docx')

print(1)

document = Document()
fake.seed(2)
random.seed(2)
for _ in range(3000):
    random.choice(func)()

document.save('no_qianzhui_b.docx')

print(1)

fake.seed(3)
random.seed(3)
for _ in range(1500):
    random.choice(func)()
document.save('no_qianzhui_c.docx')

print(1)

document = Document()
fake.seed(4)
random.seed(4)
for _ in range(3000):
    random.choice(func)()

document.save('no_qianzhui_d.docx')

print(1)

fake.seed(5)
random.seed(5)
for _ in range(1500):
    random.choice(func)()
document.save('no_qianzhui_e.docx')

print(1)

document = Document()
fake.seed(6)
random.seed(6)
for _ in range(3000):
    random.choice(func)()

document.save('no_qianzhui_f.docx')
